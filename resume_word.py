from docx import Document
from docx.shared import Inches

def get_paragraphs_text(path):
    document = Document(path)
    col_keys = []
    col_values = []
    index_num = 0
    #添加一个去重机制
    fore_str = ""
    cell_text = ""
    #提取表
    for table in document.tables:
        for row_index,row in enumerate(table.rows):
            for col_index,cell in enumerate(row.cells):
                if fore_str != cell.text:
                    if index_num % 2 ==0:
                        col_keys.append(cell.text)
                    else:
                        col_values.append(cell.text)
                    fore_str = cell.text
                    index_num += 1
                    cell_text += cell.text + '\n'
    #提取正文文本
    paragraphs_text = ""
    for paragraph in document.paragraphs:
        paragraphs_text += paragraph.text + "\n"

    return cell_text,paragraphs_text

cell_text,paragraphs_text = get_paragraphs_text(r'C:\Users\Admin\Downloads\sample\resume_sample_20200120\docx\cbb7a43eb62f.docx')
print(cell_text,paragraphs_text)